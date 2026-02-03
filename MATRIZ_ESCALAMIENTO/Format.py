from flask import Flask, render_template_string, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

app = Flask(__name__)

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Escalamiento de Atención y Soporte de Averías</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: Arial, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 20px;
            min-height: 100vh;
        }
        .container {
            max-width: 1400px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            overflow: hidden;
        }
        .header {
            background: linear-gradient(135deg, #2dd4bf 0%, #14b8a6 100%);
            color: white;
            padding: 25px;
            text-align: center;
            font-size: 24px;
            font-weight: bold;
            text-transform: uppercase;
        }
        .form-section {
            background: #f8fafc;
            padding: 30px;
            border-bottom: 3px solid #e2e8f0;
        }
        .form-section h2 {
            color: #1e293b;
            margin-bottom: 20px;
            font-size: 20px;
        }
        .form-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }
        .form-group { display: flex; flex-direction: column; }
        .form-group label {
            font-weight: 600;
            color: #475569;
            margin-bottom: 8px;
            font-size: 14px;
        }
        .form-group input {
            padding: 12px 15px;
            border: 2px solid #e2e8f0;
            border-radius: 8px;
            font-size: 15px;
            transition: all 0.3s;
        }
        .form-group input:focus {
            outline: none;
            border-color: #14b8a6;
            box-shadow: 0 0 0 3px rgba(20, 184, 166, 0.1);
        }
        .divider {
            height: 2px;
            background: linear-gradient(90deg, transparent, #14b8a6, transparent);
            margin: 20px 0;
        }
        .row-label {
            background: #fef08a;
            padding: 10px 15px;
            border-radius: 8px;
            font-weight: bold;
            color: #854d0e;
            margin-bottom: 15px;
            display: inline-block;
        }
        .download-section {
            padding: 30px;
            background: #f0fdf4;
            text-align: center;
            border-bottom: 3px solid #e2e8f0;
        }
        .download-btn {
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            padding: 15px 40px;
            font-size: 18px;
            font-weight: bold;
            border: none;
            border-radius: 10px;
            cursor: pointer;
            box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3);
            transition: all 0.3s;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        .download-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(16, 185, 129, 0.4);
        }
        .download-btn:active {
            transform: translateY(0);
        }
        .table-container { padding: 30px; overflow-x: auto; }
        table {
            width: 100%;
            border-collapse: collapse;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        th {
            background: #fef08a;
            color: #1e293b;
            padding: 15px 10px;
            text-align: center;
            font-weight: bold;
            border: 1px solid #d4d4d8;
            font-size: 13px;
        }
        td {
            padding: 12px 10px;
            border: 1px solid #d4d4d8;
            text-align: center;
            font-size: 13px;
            background: white;
        }
        td:first-child {
            background: #fef08a;
            font-weight: bold;
            color: #854d0e;
        }
        tr:hover td { background: #f1f5f9; }
        tr:hover td:first-child { background: #fde047; }
        .editable-cell {
            background: #f0fdf4 !important;
            font-weight: 600;
            color: #166534;
        }
        tr:hover .editable-cell { background: #dcfce7 !important; }
        .email-link {
            color: #0284c7;
            text-decoration: none;
            display: block;
            margin: 2px 0;
        }
        .email-link:hover { text-decoration: underline; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">ESCALAMIENTO DE ATENCIÓN Y SOPORTE DE AVERÍAS</div>
        
        <div class="form-section">
            <h2> Completar los datos </h2>
            
            <form id="dataForm" method="POST" action="/descargar">
                <div class="row-label"> Gestor de servicios </div>
                <div class="form-grid">
                    <div class="form-group">
                        <label>Nombre:</label>
                        <input type="text" name="nombre3" id="nombre3" placeholder="Ej: Alonso Pardo" value="" required>
                    </div>
                    <div class="form-group">
                        <label>Teléfono:</label>
                        <input type="text" name="telefono3" id="telefono3" placeholder="Ej: 939 111 626" value="" required>
                    </div>
                    <div class="form-group">
                        <label>Correo:</label>
                        <input type="email" name="correo3" id="correo3" placeholder="Ej: ejemplo@bitel.com.pe" value="" required>
                    </div>
                </div>

                <div class="divider"></div>

                <div class="row-label"> Gerente de Cuenta </div>
                <div class="form-grid">
                    <div class="form-group">
                        <label>Nombre:</label>
                        <input type="text" name="nombre4" id="nombre4" placeholder="Ej: Carlos Rojas" value="" required>
                    </div>
                    <div class="form-group">
                        <label>Teléfono:</label>
                        <input type="text" name="telefono4" id="telefono4" placeholder="Ej: 930 935 871" value="" required>
                    </div>
                    <div class="form-group">
                        <label>Correo:</label>
                        <input type="email" name="correo4" id="correo4" placeholder="Ej: ejemplo@bitel.com.pe" value="" required>
                    </div>
                </div>
            </form>
        </div>

        <div class="download-section">
            <button class="download-btn" onclick="document.getElementById('dataForm').submit()">
                DESCARGAR EN WORD
            </button>
        </div>

        <div class="table-container">
            <table>
                <thead>
                    <tr>
                        <th>Nivel</th>
                        <th>Tiempo transcurrido<br>desde ocurrido el<br>incidente</th>
                        <th>Departamento<br>Responsable</th>
                        <th>Cargo</th>
                        <th>Nombre</th>
                        <th>Teléfono</th>
                        <th>Correo</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td>1</td><td>Inmediato</td><td>NOC</td>
                        <td>Personal de Mesa<br>de Ayuda</td>
                        <td>Responsable en Turno</td><td><strong>Call center:</strong><br><br>0 800 79079</td>
                        <td>
                            <strong>Dirigido:</strong><br>
                            <a href="mailto:atc.corp@bitel.com.pe" class="email-link">atc.corp@bitel.com.pe</a>
                            <br>
                            <strong>Copia:</strong><br>
                            <a href="mailto:nocservicedesk@bitel.com.pe" class="email-link">nocservicedesk@bitel.com.pe</a>
                            <a href="mailto:servicemanager.corp@bitel.com.pe" class="email-link">servicemanager.corp@bitel.com.pe</a>
                        </td>
                    </tr>
                    <tr>
                        <td>2</td><td>Inmediato</td><td>NOC</td><td>Líder de Turno</td>
                        <td>Responsable en Turno<br></td>
                        <td>930 989 898</td>
                        <td><a href="mailto:nocperu@viettel.com.vn" class="email-link">nocperu@viettel.com.vn</a></td>
                    </tr>
                    <tr>
                        <td>3</td><td>1hrs-2hrs</td><td>CORPORATIVO</td><td>Gestor de servicios</td>
                        <td class="editable-cell" id="display-nombre3">Alonso Pardo</td>
                        <td class="editable-cell" id="display-telefono3">939 111 626</td>
                        <td class="editable-cell"><a href="#" id="display-correo3-link" class="email-link">pardolanamachia@bitel.com.pe</a></td>
                    </tr>
                    <tr>
                        <td>4</td><td>2hrs-3hrs</td><td>CORPORATIVO</td><td>Gerente de Cuenta</td>
                        <td class="editable-cell" id="display-nombre4">Carlos Rojas</td>
                        <td class="editable-cell" id="display-telefono4">930 935 871</td>
                        <td class="editable-cell"><a href="#" id="display-correo4-link" class="email-link">carlos.rojas@bitel.com.pe</a></td>
                    </tr>
                    <tr>
                        <td>5</td><td>3hrs-4hrs</td><td>CORPORATIVO</td><td>Director Comercial</td>
                        <td>Dennis Laime</td><td>930 800 080</td>
                        <td><a href="mailto:dennis.laime@bitel.com.pe" class="email-link">dennis.laime@bitel.com.pe</a></td>
                    </tr>
                </tbody>
            </table>
        </div>
    </div>

    <script>
        function actualizarTabla() {
            document.getElementById('display-nombre3').textContent = document.getElementById('nombre3').value || '-';
            document.getElementById('display-telefono3').textContent = document.getElementById('telefono3').value || '-';
            const correo3 = document.getElementById('correo3').value || '';
            const link3 = document.getElementById('display-correo3-link');
            link3.textContent = correo3 || '-';
            link3.href = correo3 ? 'mailto:' + correo3 : '#';
            
            document.getElementById('display-nombre4').textContent = document.getElementById('nombre4').value || '-';
            document.getElementById('display-telefono4').textContent = document.getElementById('telefono4').value || '-';
            const correo4 = document.getElementById('correo4').value || '';
            const link4 = document.getElementById('display-correo4-link');
            link4.textContent = correo4 || '-';
            link4.href = correo4 ? 'mailto:' + correo4 : '#';
        }

        ['nombre3', 'telefono3', 'correo3', 'nombre4', 'telefono4', 'correo4'].forEach(id => {
            document.getElementById(id).addEventListener('input', actualizarTabla);
        });

        actualizarTabla();
    </script>
</body>
</html>
'''

def crear_documento_word(datos):
    """Crear el documento Word """
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import RGBColor
    from docx.enum.section import WD_ORIENT
    
    doc = Document()
    
    # CONFIGURAR PÁGINA HORIZONTAL (LANDSCAPE)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Intercambiar ancho y alto para orientación horizontal
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    # Márgenes más pequeños para aprovechar el espacio
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Función para establecer color de fondo
    def set_cell_background(cell, fill_color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Crear tabla (7 filas: 1 título + 1 encabezado + 5 datos)
    table = doc.add_table(rows=7, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # FILA 1: TÍTULO COMBINADO CON FONDO TURQUESA
    # Combinar todas las celdas de la primera fila
    celda_titulo = table.rows[0].cells[0]
    for i in range(1, 7):
        celda_titulo.merge(table.rows[0].cells[i])
    
    # Agregar el título
    celda_titulo.text = 'ESCALAMIENTO DE ATENCIÓN Y SOPORTE DE AVERÍAS'
    celda_titulo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Formato del título
    for paragraph in celda_titulo.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Fondo turquesa para el título
    set_cell_background(celda_titulo, '14B8A6')
    
    # FILA 2: ENCABEZADOS con fondo amarillo
    encabezados = [
        'Nivel',
        'Tiempo\ntranscurrido\ndesde ocurrido el\nincidente',
        'Departamento\nResponsable',
        'Cargo',
        'Nombre',
        'Teléfono',
        'Correo'
    ]
    
    for i, texto in enumerate(encabezados):
        celda = table.rows[1].cells[i]
        celda.text = texto
        celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Formato del texto
        for paragraph in celda.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
        
        # Fondo amarillo
        set_cell_background(celda, 'FFFF00')
    
    # Datos de las filas (ahora empiezan desde la fila 2)
    filas_datos = [
        ['1', 'Inmediato', 'NOC', 'Personal de Mesa de Ayuda', 'Responsable en Turno', '0 800 79079', 
         'atc.corp@bitel.com.pe\nnocservicedesk@bitel.com.pe\nservicemanager.corp@bitel.com.pe'],
        
        ['2', 'Inmediato', 'NOC', 'Líder de Turno', 'Responsable en Turno\nResidente:\nMiguel Bimadi', 
         '905 785 250', 'nocperu@viettel.com.vn'],
        
        ['3', '1hrs-2hrs', 'CORPORATIVO', 'Gestor de servicios', datos['nombre3'], datos['telefono3'], datos['correo3']],
        
        ['4', '2hrs-3hrs', 'CORPORATIVO', 'Gerente de Cuenta', datos['nombre4'], datos['telefono4'], datos['correo4']],
        
        ['5', '3hrs-4hrs', 'CORPORATIVO', 'Director Comercial', 'Dennis Laime', '930 800 080', 'dennis.laime@bitel.com.pe']
    ]
    
    # Rellenar filas de datos (empezando desde fila 2, índice 2)
    for fila_idx, fila_datos in enumerate(filas_datos, start=2):
        for col_idx, valor in enumerate(fila_datos):
            celda = table.rows[fila_idx].cells[col_idx]
            celda.text = valor
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Formato del texto
            for paragraph in celda.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    
                    # Correos en azul
                    if col_idx == 6 and '@' in valor:
                        run.font.color.rgb = RGBColor(0, 0, 255)
            
            # Fondo amarillo para la columna "Nivel"
            if col_idx == 0:
                set_cell_background(celda, 'FFFF00')
                for paragraph in celda.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Ajustar anchos de columnas (más amplios por orientación horizontal)
    widths = [0.8, 1.8, 1.6, 1.8, 1.8, 1.4, 2.5]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    
    # Aplicar bordes a toda la tabla
    def set_table_borders(table):
        tbl = table._element
        tblPr = tbl.xpath('./w:tblPr')[0]
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    set_table_borders(table)
    
    return doc

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/descargar', methods=['POST'])
def descargar():
    # Obtener datos del formulario
    datos = {
        'nombre3': request.form.get('nombre3', ''),
        'telefono3': request.form.get('telefono3', ''),
        'correo3': request.form.get('correo3', ''),
        'nombre4': request.form.get('nombre4', ''),
        'telefono4': request.form.get('telefono4', ''),
        'correo4': request.form.get('correo4', '')
    }
    
    # Crear documento
    doc = crear_documento_word(datos)
    
    # Guardar en memoria
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Escalamiento_Atencion_Soporte.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    print("=" * 60)
    print("SERVIDOR FLASK INICIADO")
    print("=" * 60)
    print("Abre tu navegador en: http://localhost:5000")
    print("Completa los datos y haz clic en 'DESCARGAR EN WORD'")
    print("Para detener: Presiona Ctrl+C")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=5000)