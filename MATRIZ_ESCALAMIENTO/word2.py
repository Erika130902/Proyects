from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def crear_documento_word2(datos):
    """Crear el documento Word con residentes dinámicos"""
    
    doc2 = Document()
    
    # CONFIGURAR PÁGINA HORIZONTAL (LANDSCAPE)
    section = doc2.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Función para establecer color de fondo
    def set_cell_background(cell, fill_color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Crear tabla (7 filas: 1 título + 1 encabezado + 5 datos)
    table = doc2.add_table(rows=7, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # FILA 1: TÍTULO COMBINADO CON FONDO TURQUESA
    celda_titulo = table.rows[0].cells[0]
    for i in range(1, 7):
        celda_titulo.merge(table.rows[0].cells[i])
    
    celda_titulo.text = 'ESCALAMIENTO DE ATENCIÓN Y SOPORTE DE AVERÍAS'
    celda_titulo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    for paragraph in celda_titulo.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    set_cell_background(celda_titulo, '14B8A6')
    
    # FILA 2: ENCABEZADOS con fondo amarillo
    encabezados = [
        'Nivel',
        'Tiempo\ntranscurrido\ndesde ocurrido el\nincidente',
        'Departamento\nResponsable',
        'Cargo',
        'Nombre',
        'Teléfono',
        'Correo'
    ]
    
    for i, texto in enumerate(encabezados):
        celda = table.rows[1].cells[i]
        celda.text = texto
        celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        for paragraph in celda.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
        
        set_cell_background(celda, 'FFFF00')
    
    # ========== CONSTRUIR TELÉFONOS Y CORREOS DE FILA 1 (CON RESIDENTES) ==========
    
    # Construir texto de teléfonos para Fila 1
    telefono_fila1 = "Call Center:\n0 800 79079\n\nResidentes:\n"
    for i, (nombre, telefono) in enumerate(zip(datos['residentes_nombres'], datos['residentes_telefonos'])):
        telefono_fila1 += f"{nombre}\n{telefono}"
        if i < len(datos['residentes_nombres']) - 1:
            telefono_fila1 += "\n\n"
    
    # Construir texto de correos para Fila 1
    correo_fila1 = f"Dirigido:\natc.corp@bitel.com.pe\n{datos['correo_residentes']}\n\nCopia:\nnocservicedesk@bitel.com.pe\nservicemanager.corp@bitel.com.pe"
    
    # ========== DATOS DE LAS FILAS ==========
    filas_datos = [
        ['1', 'Inmediato', 'ATC', 'Responsable Mesa de Servicios', 'Responsable en Turno', 
         telefono_fila1, correo_fila1],
        
        ['2', 'Inmediato', 'NOC', 'Líder de Turno', 'Responsable en Turno', 
         '930 989 898', 'nocperu@viettel.com.vn'],
        
        ['3', '1hrs', 'CORPORATIVO', 'Gestor de Servicios', datos['nombre2'], datos['telefono2'], datos['correo2']],
        
        ['4', '2hrs', 'CORPORATIVO', 'Gerente de Cuenta', datos['nombre3'], datos['telefono3'], datos['correo3']],
        
        ['5', '3hrs', 'CORPORATIVO', 'Vice Director', 'Dennis Laime', '930 800 080', 'dennis.laime@bitel.com.pe']
    ]
    
    # Rellenar filas de datos (empezando desde fila 3, índice 3)
    for fila_idx, fila_datos in enumerate(filas_datos, start=2):
        for col_idx, valor in enumerate(fila_datos):
            celda = table.rows[fila_idx].cells[col_idx]
            celda.text = valor
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            for paragraph in celda.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    
                    # Correos en azul
                    if col_idx == 6 and '@' in valor:
                        run.font.color.rgb = RGBColor(0, 0, 255)
            
            # Fondo amarillo para la columna "Nivel"
            if col_idx == 0:
                set_cell_background(celda, 'FFFF00')
                for paragraph in celda.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Ajustar anchos de columnas (más amplios por orientación horizontal)
    widths = [0.8, 1.8, 1.6, 1.8, 1.8, 1.4, 2.5]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    
    # Aplicar bordes a toda la tabla
    def set_table_borders(table):
        tbl = table._element
        tblPr = tbl.xpath('./w:tblPr')[0]
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    set_table_borders(table)
    
    return doc2