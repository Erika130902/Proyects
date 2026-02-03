from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def crear_documento_word(datos):
    """Crear el documento Word """
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import RGBColor
    from docx.enum.section import WD_ORIENT
    
    doc = Document()
    
    # CONFIGURAR PÁGINA HORIZONTAL (LANDSCAPE)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Intercambiar ancho y alto para orientación horizontal
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    # Márgenes más pequeños para aprovechar el espacio
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Función para establecer color de fondo
    def set_cell_background(cell, fill_color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Crear tabla (7 filas: 1 título + 1 encabezado + 5 datos)
    table = doc.add_table(rows=7, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # FILA 1: TÍTULO COMBINADO CON FONDO TURQUESA
    # Combinar todas las celdas de la primera fila
    celda_titulo = table.rows[0].cells[0]
    for i in range(1, 7):
        celda_titulo.merge(table.rows[0].cells[i])
    
    # Agregar el título
    celda_titulo.text = 'ESCALAMIENTO DE ATENCIÓN Y SOPORTE DE AVERÍAS'
    celda_titulo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Formato del título
    for paragraph in celda_titulo.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Fondo turquesa para el título
    set_cell_background(celda_titulo, '14B8A6')
    
    # FILA 2: ENCABEZADOS con fondo amarillo
    encabezados = [
        'Nivel',
        'Tiempo\ntranscurrido\ndesde ocurrido el\nincidente',
        'Departamento\nResponsable',
        'Cargo',
        'Nombre',
        'Teléfono',
        'Correo'
    ]
    
    for i, texto in enumerate(encabezados):
        celda = table.rows[1].cells[i]
        celda.text = texto
        celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Formato del texto
        for paragraph in celda.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
        
        # Fondo amarillo
        set_cell_background(celda, 'FFFF00')
    
    # Datos de las filas (ahora empiezan desde la fila 2)
    filas_datos = [
        ['1', 'Inmediato', 'NOC', 'Personal de Mesa de Ayuda', 'Responsable en Turno', '0 800 79079', 
         'Dirigido:\natc.corp@bitel.com.pe\nCopia:\nnocservicedesk@bitel.com.pe\nservicemanager.corp@bitel.com.pe'],
        
        ['2', 'Inmediato', 'NOC', 'Líder de Turno', 'Responsable en Turno', 
         'Call center:\n930 989 898', 'nocperu@viettel.com.vn'],
        
        ['3', '1hrs-2hrs', 'CORPORATIVO', 'Gestor de servicios', datos['nombre3'], datos['telefono3'], datos['correo3']],
        
        ['4', '2hrs-3hrs', 'CORPORATIVO', 'Gerente de Cuenta', datos['nombre4'], datos['telefono4'], datos['correo4']],
        
        ['5', '3hrs-4hrs', 'CORPORATIVO', 'Director Comercial', 'Dennis Laime', '930 800 080', 'dennis.laime@bitel.com.pe']
    ]
    
    # Rellenar filas de datos (empezando desde fila 2, índice 2)
    for fila_idx, fila_datos in enumerate(filas_datos, start=2):
        for col_idx, valor in enumerate(fila_datos):
            celda = table.rows[fila_idx].cells[col_idx]
            celda.text = valor
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Formato del texto
            for paragraph in celda.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    
                    # Correos en azul
                    if col_idx == 6 and '@' in valor:
                        run.font.color.rgb = RGBColor(0, 0, 255)
            
            # Fondo amarillo para la columna "Nivel"
            if col_idx == 0:
                set_cell_background(celda, 'FFFF00')
                for paragraph in celda.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Ajustar anchos de columnas (más amplios por orientación horizontal)
    widths = [0.8, 1.8, 1.6, 1.8, 1.8, 1.4, 2.5]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    
    # Aplicar bordes a toda la tabla
    def set_table_borders(table):
        tbl = table._element
        tblPr = tbl.xpath('./w:tblPr')[0]
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    set_table_borders(table)
    
    return doc